"""
Resume Parser Service
=====================
Extracts text and metadata from PDF and DOCX resume files.
Uses pdfplumber (primary) with PyPDF2 as fallback for PDFs.
Uses python-docx for DOCX files.
"""

import io
import re
from typing import List, Optional


class ResumeParser:
    def extract_from_pdf(self, content: bytes) -> str:
        # pdf extraction logic
        return "extracted text"

    def extract_skills(self, text: str) -> List[str]:
        # NLP logic
        return ["Python", "FastAPI"]

    # ... other methods (extract_education, etc.)

# PDF parsing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# NLP
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except Exception:
    SPACY_AVAILABLE = False


# ── Skill Master List (extend this per industry) ──────────────────────────────
SKILL_KEYWORDS = {
    # Programming languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab",
    # Web / Frontend
    "react", "vue", "angular", "next.js", "nuxt", "svelte", "html", "css",
    "tailwind", "bootstrap", "webpack", "vite",
    # Backend / Infra
    "node.js", "fastapi", "django", "flask", "express", "spring", "graphql",
    "rest api", "microservices", "docker", "kubernetes", "aws", "gcp", "azure",
    "terraform", "ci/cd", "jenkins", "github actions",
    # Data / AI
    "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "llm",
    "machine learning", "deep learning", "nlp", "data analysis", "tableau", "power bi",
    # Soft / PM
    "agile", "scrum", "jira", "confluence", "product management", "stakeholder management",
    "communication", "leadership", "project management", "cross-functional",
}

SECTION_HEADERS = [
    "experience", "work experience", "employment", "professional experience",
    "education", "academic background",
    "skills", "technical skills", "core competencies",
    "summary", "objective", "profile", "about",
    "projects", "certifications", "awards", "publications",
    "volunteer", "languages",
]


class ResumeParser:

    # ── PDF ───────────────────────────────────────────────────────────────────

    def extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF bytes. Uses pdfplumber first, PyPDF2 as fallback."""
        text = ""

        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text = "\n".join(
                        page.extract_text() or "" for page in pdf.pages
                    )
            except Exception:
                pass

        if not text.strip() and PYPDF2_AVAILABLE:
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
            except Exception:
                pass

        if not text.strip():
            raise ValueError("Could not extract text from PDF. The file may be scanned or image-based.")

        return self._clean_text(text)

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX bytes using python-docx."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        return self._clean_text("\n".join(paragraphs))

    # ── NLP Extraction ────────────────────────────────────────────────────────

    def extract_skills(self, text: str) -> List[str]:
        """Extract skills using keyword matching + optional spaCy NER."""
        text_lower = text.lower()
        found = [s for s in SKILL_KEYWORDS if s in text_lower]

        if SPACY_AVAILABLE:
            doc = nlp(text[:5000])  # limit for performance
            for ent in doc.ents:
                if ent.label_ in ("ORG", "PRODUCT"):
                    candidate = ent.text.lower().strip()
                    if len(candidate) > 1:
                        found.append(candidate)

        return sorted(set(found))

    def extract_experience_years(self, text: str) -> Optional[int]:
        """Estimate total years of experience from date ranges in resume."""
        # Pattern: 2018 - 2021, Jan 2019 – Present, etc.
        year_pattern = re.findall(r"\b(19|20)\d{2}\b", text)
        if len(year_pattern) >= 2:
            years = [int(y) for y in year_pattern]
            return max(years) - min(years)
        return None

    def extract_education(self, text: str) -> List[str]:
        """Extract degree-level statements."""
        patterns = [
            r"(bachelor|master|ph\.?d|b\.?s|m\.?s|mba|b\.?e|m\.?tech|b\.?tech)[^\n,;]*",
        ]
        results = []
        for pat in patterns:
            results.extend(re.findall(pat, text, re.IGNORECASE))
        return list(set(r.strip() for r in results if r.strip()))

    def detect_sections(self, text: str) -> List[str]:
        """Detect which standard resume sections are present."""
        text_lower = text.lower()
        return [s for s in SECTION_HEADERS if s in text_lower]

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        """Normalize whitespace and remove junk characters."""
        text = re.sub(r"\x00", "", text)          # null bytes
        text = re.sub(r"\r\n|\r", "\n", text)     # normalize newlines
        text = re.sub(r"\n{3,}", "\n\n", text)    # collapse blank lines
        text = re.sub(r" {2,}", " ", text)         # collapse spaces
        return text.strip()
